"""
EXCEL_interfaz_libro_capitulos
"""

#!/usr/bin/env python
# coding: utf-8

# In[2]:


from langchain.prompts import PromptTemplate
from langchain_openai.chat_models import ChatOpenAI  # Actualizado para usar el nuevo paquete
import os
from docx import Document
from dotenv import load_dotenv

# Cargar las variables de entorno desde el archivo .env
load_dotenv()

# Obtener la API key de las variables de entorno
api_key = os.getenv("OPENAI_API_KEY")

# Verificar si la API key se cargó correctamente
if not api_key:
    raise ValueError("La clave de la API de OpenAI no se encontró en las variables de entorno.")

# Crear el cliente de LangChain usando la API key cargada
llm = ChatOpenAI(api_key=api_key, model="o3-mini-2025-01-31", temperature= 1)

# Definir plantillas de prompt
prompt_template_1 = PromptTemplate(
    input_variables=["capitulo", "pestania", "grupo", "funcion"],
    template="""Siguiendo las siguientes reglas:
1.	El largo de las oraciones debe ser mayor que el normal.
2.	Utiliza menos puntos separadores de oraciones y conecta las oraciones utilizando palabras conectoras del idioma español.
3.	Evita la redundancia y la divagación: Expón las ideas de manera directa y sin desviaciones innecesarias.
4.	Usa un lenguaje accesible: Prioriza palabras comunes sobre términos innecesariamente complejos.
5.	Adapta el nivel de lectura a una audiencia universitaria: Usa un lenguaje preciso, sin tecnicismos innecesarios que dificulten la comprensión.
6.	Las frases de tus repuestas deben contener dos oraciones largas
7.	Conecta las ideas de manera natural: Evita transiciones forzadas o artificiosas entre párrafos y secciones.
8.	Sigue las normas convencionales de puntuación: Usa signos de puntuación estándar y evita abusar de los mismos.
9.	Varía la estructura de las oraciones: Alterna la construcción sintáctica para evitar la monotonía en el discurso.
10.	La respuesta debe ser una sola frase con varias oraciones
11.	Evita el uso de conectores innecesarios: Reduce la presencia de términos como "de hecho", "además", "por lo tanto", "asimismo", "sin embargo", "aparentemente", "en consecuencia", "específicamente", "notablemente" y "alternativamente", ya que pueden sobrecargar el texto y restar claridad.
12.	Asegúrate de que las oraciones se mantengan coherentes y fáciles de entender a pesar de su longitud.
13.	Considera el tono y el contexto original del texto al hacer las transformaciones.
Output Format
•	El texto debe ser fluido, imitando el estilo natural y conversacional de un hablante nativo del español, sin perder el significado original. Utilizando menos puntos separadores de oraciones y más palabras de conexión del español
 Notes

Esta es la lista de conectores que debes usar entre oraciones :
Aquí tienes una lista de conectores en español, organizados por categorías:
________________________________________
CONECTORES ADITIVOS (Añaden información)
•	además
•	agregado a lo anterior
•	asimismo
•	de igual forma
•	de igual manera
•	de la misma manera
•	del mismo modo
•	en la misma línea
•	encima
•	es más
•	hasta
•	igualmente
•	incluso
•	más aún
•	para colmo
•	por añadidura
•	también
•	y
________________________________________
CONECTORES DE CONTRASTE Y OPOSICIÓN
Conectores concesivos (Muestran dificultad o contradicción)
•	a pesar de que
•	a pesar de todo
•	ahora bien
•	al mismo tiempo
•	aun así
•	aun cuando
•	aunque
•	con todo
•	de cualquier modo
Conectores restrictivos (Limitan o contrastan una idea)
•	al contrario
•	en cambio
•	en cierta medida
•	en cierto modo
•	hasta cierto punto
•	no obstante
•	pero
•	por otra parte
•	sin embargo
________________________________________
CONECTORES DE CAUSA (Explican razones o motivos)
•	dado que
•	debido a que
•	porque
•	pues
•	puesto que
•	ya que
________________________________________
CONECTORES DE CONSECUENCIA (Expresan efectos o resultados)
•	así pues
•	así que
•	de ahí que
•	de manera que
•	de tal forma
•	en consecuencia
•	en ese sentido
•	entonces
•	luego
•	por consiguiente
•	por eso
•	por esta razón
•	por lo que sigue
•	por lo tanto
•	por tanto
________________________________________
CONECTORES COMPARATIVOS (Establecen semejanzas o diferencias)
•	análogamente
•	así como
•	como
•	de modo similar
•	del mismo modo
•	igual… que…
•	igualmente
•	más/menos… que…
•	tan… como…
________________________________________
CONECTORES REFORMULATIVOS
Conectores de explicación (Aclaran o explican una idea)
•	a saber
•	dicho de otro modo
•	en otras palabras
•	es decir
•	esto es
•	o sea
Conectores de recapitulación (Resumen lo dicho anteriormente)
•	dicho de otro modo
•	en breve
•	en otras palabras
•	en resumen
•	en resumidas cuentas
•	en síntesis
•	en suma
•	en una palabra
•	total
Conectores de ejemplificación (Introducen ejemplos)
•	así como
•	específicamente
•	para ilustrar
•	particularmente
•	por ejemplo
•	tal es el caso de
Conectores de corrección (Corrigen o reformulan una idea)
•	a decir verdad
•	mejor dicho
•	o sea
________________________________________
CONECTORES ORDENADORES
Conectores para iniciar el discurso
•	ante todo
•	bueno
•	en primer lugar
•	en principio
•	para comenzar
•	primeramente
Conectores para cerrar el discurso
•	al final
•	en conclusión
•	en fin
•	en suma
•	finalmente
•	para concluir
•	para finalizar
•	para resumir
•	por último
Conectores de transición
•	a continuación
•	acto seguido
•	ahora bien
•	después / luego
•	por otra parte
•	por otro lado
Conectores de digresión (Introducen una idea secundaria)
•	a propósito
•	a todo esto
•	con respecto a
•	en cuanto a
•	por cierto
•	por otra parte
________________________________________
CONECTORES TEMPORALES (Ubican un evento en el tiempo)
•	a partir de
•	actualmente
•	al principio
•	antes de
•	apenas
•	cuando
•	desde (entonces)
•	después de
•	en cuanto
•	en el comienzo
•	hasta que
•	inmediatamente
•	luego
•	no bien
•	temporalmente
________________________________________
CONECTORES ESPACIALES (Ubican un evento en el espacio)
•	a la izquierda/derecha
•	abajo
•	al lado
•	arriba
•	en el fondo/medio
•	adelante
________________________________________
CONECTORES CONDICIONALES (Expresan condiciones o requisitos)
•	a menos que
•	a no ser que
•	con tal que
•	en caso de que
•	mientras que
•	según
•	si
•	siempre que
•	siempre y cuando
________________________________________
CONECTORES DE CERTEZA (Expresan seguridad o certeza)
•	ciertamente
•	claro está
•	como es por muchos conocido
•	como nadie ignora
•	con certeza
•	con seguridad
•	efectivamente
•	en realidad
•	en verdad
•	es evidente
•	indudablemente
•	realmente
________________________________________
CONECTORES DE FINALIDAD (Expresan el propósito de una acción)
•	a fin de
•	con el objetivo de
•	con el propósito de
•	con la intención de
•	de manera que
•	de tal forma que
•	de modo que
•	para
•	con el objeto de


Basándote en las instrucciones elabora un texto corto con los siguientes detalles:
Capítulo: {capitulo}
Pestaña: {pestania}
Grupo: {grupo}
Función: {funcion}

Proporciona un resumen detallado de los siguientes aspectos de esta funcion:función {funcion} en el contexto de Excel, dentro del grupo {grupo} de la pestaña {pestania}
1.- Cual es el proposito de esta funcion. 
2.- Cual es el encaminamien que se debe serguir en excel para enciomntrar esta funcion. 
3.- Brinda un ejemplo de caso de uso. 
La explicación debe ser clara, concisa y mantener un tono profesional. Limítate a un máximo de 150 tokens.
"""
)

# prompt_template_2 = PromptTemplate(
#     input_variables=["capitulo", "pestania", "grupo", "funcion"],
#     template="""Sigue las siguientes reglas:
# 1.	El largo de las oraciones debe ser mayor que el normal.
# 2.	Utiliza menos puntos separadores de oraciones y conecta las oraciones utilizando palabras conectoras del idioma español.
# 3.	Evita la redundancia y la divagación: Expón las ideas de manera directa y sin desviaciones innecesarias.
# 4.	Usa un lenguaje accesible: Prioriza palabras comunes sobre términos innecesariamente complejos.
# 5.	Adapta el nivel de lectura a una audiencia universitaria: Usa un lenguaje preciso, sin tecnicismos innecesarios que dificulten la comprensión.
# 6.	Las frases de tus repuestas deben contener dos oraciones largas
# 7.	Conecta las ideas de manera natural: Evita transiciones forzadas o artificiosas entre párrafos y secciones.
# 8.	Sigue las normas convencionales de puntuación: Usa signos de puntuación estándar y evita abusar de los mismos.
# 9.	Varía la estructura de las oraciones: Alterna la construcción sintáctica para evitar la monotonía en el discurso.
# 10.	La respuesta debe ser una sola frase con varias oraciones
# 11.	Evita el uso de conectores innecesarios: Reduce la presencia de términos como "de hecho", "además", "por lo tanto", "asimismo", "sin embargo", "aparentemente", "en consecuencia", "específicamente", "notablemente" y "alternativamente", ya que pueden sobrecargar el texto y restar claridad.
# 12.	Asegúrate de que las oraciones se mantengan coherentes y fáciles de entender a pesar de su longitud.
# 13.	Considera el tono y el contexto original del texto al hacer las transformaciones.
# Output Format
# •	El texto debe ser fluido, imitando el estilo natural y conversacional de un hablante nativo del español, sin perder el significado original. Utilizando menos puntos separadores de oraciones y más palabras de conexión del español
#  Notes

# Esta es la lista de conectores que debes usar entre oraciones :
# Aquí tienes una lista de conectores en español, organizados por categorías:
# ________________________________________
# CONECTORES ADITIVOS (Añaden información)
# •	además
# •	agregado a lo anterior
# •	asimismo
# •	de igual forma
# •	de igual manera
# •	de la misma manera
# •	del mismo modo
# •	en la misma línea
# •	encima
# •	es más
# •	hasta
# •	igualmente
# •	incluso
# •	más aún
# •	para colmo
# •	por añadidura
# •	también
# •	y
# ________________________________________
# CONECTORES DE CONTRASTE Y OPOSICIÓN
# Conectores concesivos (Muestran dificultad o contradicción)
# •	a pesar de que
# •	a pesar de todo
# •	ahora bien
# •	al mismo tiempo
# •	aun así
# •	aun cuando
# •	aunque
# •	con todo
# •	de cualquier modo
# Conectores restrictivos (Limitan o contrastan una idea)
# •	al contrario
# •	en cambio
# •	en cierta medida
# •	en cierto modo
# •	hasta cierto punto
# •	no obstante
# •	pero
# •	por otra parte
# •	sin embargo
# ________________________________________
# CONECTORES DE CAUSA (Explican razones o motivos)
# •	dado que
# •	debido a que
# •	porque
# •	pues
# •	puesto que
# •	ya que
# ________________________________________
# CONECTORES DE CONSECUENCIA (Expresan efectos o resultados)
# •	así pues
# •	así que
# •	de ahí que
# •	de manera que
# •	de tal forma
# •	en consecuencia
# •	en ese sentido
# •	entonces
# •	luego
# •	por consiguiente
# •	por eso
# •	por esta razón
# •	por lo que sigue
# •	por lo tanto
# •	por tanto
# ________________________________________
# CONECTORES COMPARATIVOS (Establecen semejanzas o diferencias)
# •	análogamente
# •	así como
# •	como
# •	de modo similar
# •	del mismo modo
# •	igual… que…
# •	igualmente
# •	más/menos… que…
# •	tan… como…
# ________________________________________
# CONECTORES REFORMULATIVOS
# Conectores de explicación (Aclaran o explican una idea)
# •	a saber
# •	dicho de otro modo
# •	en otras palabras
# •	es decir
# •	esto es
# •	o sea
# Conectores de recapitulación (Resumen lo dicho anteriormente)
# •	dicho de otro modo
# •	en breve
# •	en otras palabras
# •	en resumen
# •	en resumidas cuentas
# •	en síntesis
# •	en suma
# •	en una palabra
# •	total
# Conectores de ejemplificación (Introducen ejemplos)
# •	así como
# •	específicamente
# •	para ilustrar
# •	particularmente
# •	por ejemplo
# •	tal es el caso de
# Conectores de corrección (Corrigen o reformulan una idea)
# •	a decir verdad
# •	mejor dicho
# •	o sea
# ________________________________________
# CONECTORES ORDENADORES
# Conectores para iniciar el discurso
# •	ante todo
# •	bueno
# •	en primer lugar
# •	en principio
# •	para comenzar
# •	primeramente
# Conectores para cerrar el discurso
# •	al final
# •	en conclusión
# •	en fin
# •	en suma
# •	finalmente
# •	para concluir
# •	para finalizar
# •	para resumir
# •	por último
# Conectores de transición
# •	a continuación
# •	acto seguido
# •	ahora bien
# •	después / luego
# •	por otra parte
# •	por otro lado
# Conectores de digresión (Introducen una idea secundaria)
# •	a propósito
# •	a todo esto
# •	con respecto a
# •	en cuanto a
# •	por cierto
# •	por otra parte
# ________________________________________
# CONECTORES TEMPORALES (Ubican un evento en el tiempo)
# •	a partir de
# •	actualmente
# •	al principio
# •	antes de
# •	apenas
# •	cuando
# •	desde (entonces)
# •	después de
# •	en cuanto
# •	en el comienzo
# •	hasta que
# •	inmediatamente
# •	luego
# •	no bien
# •	temporalmente
# ________________________________________
# CONECTORES ESPACIALES (Ubican un evento en el espacio)
# •	a la izquierda/derecha
# •	abajo
# •	al lado
# •	arriba
# •	en el fondo/medio
# •	adelante
# ________________________________________
# CONECTORES CONDICIONALES (Expresan condiciones o requisitos)
# •	a menos que
# •	a no ser que
# •	con tal que
# •	en caso de que
# •	mientras que
# •	según
# •	si
# •	siempre que
# •	siempre y cuando
# ________________________________________
# CONECTORES DE CERTEZA (Expresan seguridad o certeza)
# •	ciertamente
# •	claro está
# •	como es por muchos conocido
# •	como nadie ignora
# •	con certeza
# •	con seguridad
# •	efectivamente
# •	en realidad
# •	en verdad
# •	es evidente
# •	indudablemente
# •	realmente
# ________________________________________
# CONECTORES DE FINALIDAD (Expresan el propósito de una acción)
# •	a fin de
# •	con el objetivo de
# •	con el propósito de
# •	con la intención de
# •	de manera que
# •	de tal forma que
# •	de modo que
# •	para
# •	con el objeto de


# Basándote en las instrucciones elabora un texto corto con los siguientes detalles:
# Capítulo: {capitulo}
# Pestaña: {pestania}
# Grupo: {grupo}
# Función: {funcion}

# Proporciona un resumen detallado del propósito de esta función, académico y enfocado en cómo se utiliza la función {funcion} en el contexto de Excel, dentro del grupo {grupo} de la pestaña {pestania}. La explicación debe ser clara, concisa y mantener un tono profesional. Limítate a un máximo de 50 tokens.
# """
# )

# prompt_template_3 = PromptTemplate(
#     input_variables=["capitulo", "pestania", "grupo", "funcion"],
#     template="""Sigue las siguientes reglas: 
# 1.	El largo de las oraciones debe ser mayor que el normal.
# 2.	Utiliza menos puntos separadores de oraciones y conecta las oraciones utilizando palabras conectoras del idioma español.
# 3.	Evita la redundancia y la divagación: Expón las ideas de manera directa y sin desviaciones innecesarias.
# 4.	Usa un lenguaje accesible: Prioriza palabras comunes sobre términos innecesariamente complejos.
# 5.	Adapta el nivel de lectura a una audiencia universitaria: Usa un lenguaje preciso, sin tecnicismos innecesarios que dificulten la comprensión.
# 6.	Las frases de tus repuestas deben contener dos oraciones largas
# 7.	Conecta las ideas de manera natural: Evita transiciones forzadas o artificiosas entre párrafos y secciones.
# 8.	Sigue las normas convencionales de puntuación: Usa signos de puntuación estándar y evita abusar de los mismos.
# 9.	Varía la estructura de las oraciones: Alterna la construcción sintáctica para evitar la monotonía en el discurso.
# 10.	La respuesta debe ser una sola frase con varias oraciones
# 11.	Evita el uso de conectores innecesarios: Reduce la presencia de términos como "de hecho", "además", "por lo tanto", "asimismo", "sin embargo", "aparentemente", "en consecuencia", "específicamente", "notablemente" y "alternativamente", ya que pueden sobrecargar el texto y restar claridad.
# 12.	Asegúrate de que las oraciones se mantengan coherentes y fáciles de entender a pesar de su longitud.
# 13.	Considera el tono y el contexto original del texto al hacer las transformaciones.
# Output Format
# •	El texto debe ser fluido, imitando el estilo natural y conversacional de un hablante nativo del español, sin perder el significado original. Utilizando menos puntos separadores de oraciones y más palabras de conexión del español
#  Notes

# Esta es la lista de conectores que debes usar entre oraciones :
# Aquí tienes una lista de conectores en español, organizados por categorías:
# ________________________________________
# CONECTORES ADITIVOS (Añaden información)
# •	además
# •	agregado a lo anterior
# •	asimismo
# •	de igual forma
# •	de igual manera
# •	de la misma manera
# •	del mismo modo
# •	en la misma línea
# •	encima
# •	es más
# •	hasta
# •	igualmente
# •	incluso
# •	más aún
# •	para colmo
# •	por añadidura
# •	también
# •	y
# ________________________________________
# CONECTORES DE CONTRASTE Y OPOSICIÓN
# Conectores concesivos (Muestran dificultad o contradicción)
# •	a pesar de que
# •	a pesar de todo
# •	ahora bien
# •	al mismo tiempo
# •	aun así
# •	aun cuando
# •	aunque
# •	con todo
# •	de cualquier modo
# Conectores restrictivos (Limitan o contrastan una idea)
# •	al contrario
# •	en cambio
# •	en cierta medida
# •	en cierto modo
# •	hasta cierto punto
# •	no obstante
# •	pero
# •	por otra parte
# •	sin embargo
# ________________________________________
# CONECTORES DE CAUSA (Explican razones o motivos)
# •	dado que
# •	debido a que
# •	porque
# •	pues
# •	puesto que
# •	ya que
# ________________________________________
# CONECTORES DE CONSECUENCIA (Expresan efectos o resultados)
# •	así pues
# •	así que
# •	de ahí que
# •	de manera que
# •	de tal forma
# •	en consecuencia
# •	en ese sentido
# •	entonces
# •	luego
# •	por consiguiente
# •	por eso
# •	por esta razón
# •	por lo que sigue
# •	por lo tanto
# •	por tanto
# ________________________________________
# CONECTORES COMPARATIVOS (Establecen semejanzas o diferencias)
# •	análogamente
# •	así como
# •	como
# •	de modo similar
# •	del mismo modo
# •	igual… que…
# •	igualmente
# •	más/menos… que…
# •	tan… como…
# ________________________________________
# CONECTORES REFORMULATIVOS
# Conectores de explicación (Aclaran o explican una idea)
# •	a saber
# •	dicho de otro modo
# •	en otras palabras
# •	es decir
# •	esto es
# •	o sea
# Conectores de recapitulación (Resumen lo dicho anteriormente)
# •	dicho de otro modo
# •	en breve
# •	en otras palabras
# •	en resumen
# •	en resumidas cuentas
# •	en síntesis
# •	en suma
# •	en una palabra
# •	total
# Conectores de ejemplificación (Introducen ejemplos)
# •	así como
# •	específicamente
# •	para ilustrar
# •	particularmente
# •	por ejemplo
# •	tal es el caso de
# Conectores de corrección (Corrigen o reformulan una idea)
# •	a decir verdad
# •	mejor dicho
# •	o sea
# ________________________________________
# CONECTORES ORDENADORES
# Conectores para iniciar el discurso
# •	ante todo
# •	bueno
# •	en primer lugar
# •	en principio
# •	para comenzar
# •	primeramente
# Conectores para cerrar el discurso
# •	al final
# •	en conclusión
# •	en fin
# •	en suma
# •	finalmente
# •	para concluir
# •	para finalizar
# •	para resumir
# •	por último
# Conectores de transición
# •	a continuación
# •	acto seguido
# •	ahora bien
# •	después / luego
# •	por otra parte
# •	por otro lado
# Conectores de digresión (Introducen una idea secundaria)
# •	a propósito
# •	a todo esto
# •	con respecto a
# •	en cuanto a
# •	por cierto
# •	por otra parte
# ________________________________________
# CONECTORES TEMPORALES (Ubican un evento en el tiempo)
# •	a partir de
# •	actualmente
# •	al principio
# •	antes de
# •	apenas
# •	cuando
# •	desde (entonces)
# •	después de
# •	en cuanto
# •	en el comienzo
# •	hasta que
# •	inmediatamente
# •	luego
# •	no bien
# •	temporalmente
# ________________________________________
# CONECTORES ESPACIALES (Ubican un evento en el espacio)
# •	a la izquierda/derecha
# •	abajo
# •	al lado
# •	arriba
# •	en el fondo/medio
# •	adelante
# ________________________________________
# CONECTORES CONDICIONALES (Expresan condiciones o requisitos)
# •	a menos que
# •	a no ser que
# •	con tal que
# •	en caso de que
# •	mientras que
# •	según
# •	si
# •	siempre que
# •	siempre y cuando
# ________________________________________
# CONECTORES DE CERTEZA (Expresan seguridad o certeza)
# •	ciertamente
# •	claro está
# •	como es por muchos conocido
# •	como nadie ignora
# •	con certeza
# •	con seguridad
# •	efectivamente
# •	en realidad
# •	en verdad
# •	es evidente
# •	indudablemente
# •	realmente
# ________________________________________
# CONECTORES DE FINALIDAD (Expresan el propósito de una acción)
# •	a fin de
# •	con el objetivo de
# •	con el propósito de
# •	con la intención de
# •	de manera que
# •	de tal forma que
# •	de modo que
# •	para
# •	con el objeto de


# Basándote en las instrucciones elabora un texto corto con los siguientes detalles:
# Capítulo: {capitulo}
# Pestaña: {pestania}
# Grupo: {grupo}
# Función: {funcion}

# Proporciona un resumen detallado del propósito de esta función, académico y enfocado en cómo se utiliza la función {funcion} en el contexto de Excel, dentro del grupo {grupo} de la pestaña {pestania}. La explicación debe ser clara, concisa y mantener un tono profesional. Limítate a un máximo de 50 tokens.
# """
# )

# Funciones para procesar las interacciones con LangChain
def interactuar_proposito(capitulo, pestania, grupo, funcion):
    chain = prompt_template_1 | llm
    response = chain.invoke({"capitulo": capitulo, "pestania": pestania, "grupo": grupo, "funcion": funcion})
    return response.content.strip()

# def interactuar_encaminamiento(capitulo, pestania, grupo, funcion):
#     chain = prompt_template_2 | llm
#     response = chain.invoke({"capitulo": capitulo, "pestania": pestania, "grupo": grupo, "funcion": funcion})
#     return response.content.strip()

# def interactuar_ejemplo(capitulo, pestania, grupo, funcion):
#     chain = prompt_template_3 | llm
#     response = chain.invoke({"capitulo": capitulo, "pestania": pestania, "grupo": grupo, "funcion": funcion})
#     return response.content.strip()

def interactuar_gpt(capitulo, pestania, grupo, funcion):
    proposito = interactuar_proposito(capitulo, pestania, grupo, funcion)
    # encaminamiento = interactuar_encaminamiento(capitulo, pestania, grupo, funcion)
    # ejemplo = interactuar_ejemplo(capitulo, pestania, grupo, funcion)
    # return f"{proposito}\n\n{encaminamiento}\n\n{ejemplo
    return f"{proposito}"


# Función para procesar el archivo JSON y generar explicaciones
def generar_resumenes_desde_json(json_data, path_out):
    doc = Document()
    doc.add_heading("Resumen de Funciones de Excel", level=1)

    for capitulo, pestañas in json_data.items():
        doc.add_heading(f"Capítulo: {capitulo}", level=1)
        for pestaña, grupos in pestañas.items():
            doc.add_heading(f"Pestaña: {pestaña}", level=2)
            for grupo, funciones in grupos.items():
                doc.add_heading(f"Grupo: {grupo}", level=3)
                for idx, funcion in enumerate(funciones, start=1):
                    print(f"Procesando: Capítulo {capitulo}, Pestaña {pestaña}, Grupo {grupo}, Función {funcion}")
                    resumen = interactuar_gpt(capitulo, pestaña, grupo, funcion)
                    doc.add_heading(f"{idx}. Función: {funcion}", level=4)
                    doc.add_paragraph(f"\n{resumen}")

    doc.save(path_out)
    print(f"Documento completado y guardado en: {path_out}")

# Datos de ejemplo
json_data = {
    "8": {
        "Programador": {
            "Código": [
                "Visual Basic",
                "Macros",
                "Grabar macro",
                "Usar referencias relativas"
            ],
            "Complementos": [
                "Complementos",
                "Complementos de Excel",
                "Complementos COM"
            ],
            "Controles": [
                "Insertar",
                "Propiedades",
                "Ver código",
                "Ejecutar cuadro de diálogo"
            ]
        }
    }
}

# Ruta de salida
capitulo = list(json_data.keys())[0]
path_out = f"C:\\Users\\HP\\Desktop\\LIBROS PERSO\\EXCEL INTERFAZ GRAFICA\\libro_excel_CAPITULO_{capitulo}.docx"
generar_resumenes_desde_json(json_data, path_out)


# In[ ]:






if __name__ == "__main__":
    pass
